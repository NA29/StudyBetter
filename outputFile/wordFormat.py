import docx
import requests
from docx.shared import Inches
import subprocess
import sys
import os
from openai import OpenAI

current_script_path = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_script_path)
sys.path.append(parent_dir)

import googleimage
from colordetect import colors
from google_images_search import GoogleImagesSearch

PATH_TXT = "output.txt"
notesTaken = ""





def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()


def getDefinition(word):
    url = f"https://api.dictionaryapi.dev/api/v2/entries/en/{word}"
    response = requests.get(url)

    if response.status_code == 200:
        data = response.json()

        if not data:
            return "No definitions found."
        
        first_definition = data[0]["meanings"][0]["definitions"][0]["definition"]
        return first_definition
    else: return "Word not found or an error occurred."

# def getQuestion(word):
#     completion = client.chat.completions.create(
#         model="gpt-3.5-turbo",
#         messages=[
#             {"role": "system", "content": f"Create a question to test someone's abilities on the subject regarding this word: {word}"},
#         ]  
#     )
#     return completion


def main_word():
    document = docx.Document()
    table = document.add_table(rows=1, cols=2)
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    with open(PATH_TXT, 'r') as file:
        content = file.read()

    # Remove \n completely
    content_without_newlines = content.replace('/n', '')
    left_cell.text = content_without_newlines
    
    dict_words = colors.main()
    print(str(dict_words) + " yes")
    for k,v in dict_words.items():
        if v == "red":
            definition = getDefinition(k.lower())
            document.add_paragraph(f"{k.lower()}: {definition}")
            print(definition)
        elif v == "green":
            #googleimage.getImage(k.lower())
            pass
            
             
            
    directory = os.fsencode("./images")          
    for file in os.listdir(directory):         
        filename = os.fsdecode(file)         
        if filename.endswith(".jpg"):             
            path_current_image = os.path.join("./images", filename)
            right_cell.add_paragraph().add_run().add_picture(path_current_image,width=Inches(2))
            continue        
        else:             
            continue
    
            

    document.save('notes.docx')
    if sys.platform.startswith('darwin'):  
        subprocess.run(['open', 'notes.docx'])
    #print(getQuestion('mitochondria'))
    


if __name__ == '__main__':
    main_word()